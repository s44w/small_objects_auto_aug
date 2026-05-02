"""
inject-titlepage.py — генерация titlepage-filled.docx из метаданных студента.

MVP-stub: на текущем этапе создаёт минимальный титульник через python-docx,
используя стили из blank-template.docx (+Тит_Абзац по центру, +Тит_Тема ВКР).
Полная вёрстка как у Маршуниной — задача следующей итерации.

Использование:
    python inject-titlepage.py \\
        --template template/blank-template.docx \\
        --output template/titlepage-filled.docx \\
        --topic "..." \\
        --student "..." \\
        --group "..." \\
        --direction "..." \\
        --advisor "..." \\
        --institute "..." \\
        --department "..." \\
        --year "..."
"""
from __future__ import annotations
import argparse
import sys
from pathlib import Path


def build_titlepage(args: argparse.Namespace) -> None:
    try:
        from docx import Document
    except ImportError:
        print("ОШИБКА: установите python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    template_path = Path(args.template).resolve()
    if not template_path.is_file():
        print(f"ОШИБКА: шаблон не найден: {template_path}", file=sys.stderr)
        sys.exit(1)

    # Открываем шаблон чтобы унаследовать все стили
    doc = Document(str(template_path))

    # Очистить тело документа (оставить только sectPr)
    body = doc.element.body
    sectPr = None
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            sectPr = child
        else:
            body.remove(child)

    style_centered = "+Тит_Абзац по центру"
    style_topic = "+Тит_Тема ВКР"

    def add(text: str, style: str = style_centered) -> None:
        p = doc.add_paragraph(text)
        try:
            p.style = doc.styles[style]
        except KeyError:
            # Стиль не найден — оставляем default Normal
            pass

    # Шапка титула
    add("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ")
    add("РОССИЙСКОЙ ФЕДЕРАЦИИ")
    add("")
    add("федеральное государственное автономное")
    add("образовательное учреждение высшего образования")
    add(f"«{args.university}»" if args.university else "«Самарский национальный исследовательский университет имени академика С.П. Королёва»")
    add("")
    add(args.institute)
    add(args.department)
    add("")
    add("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА")
    add("")
    # Тема — отдельным стилем
    add(args.topic, style=style_topic)
    add("")
    add(f"по программе бакалавриата по направлению подготовки")
    add(args.direction)
    add("")
    add("")
    # Подписи
    p = doc.add_paragraph()
    p.add_run(f"Обучающийся         {args.student}    ___________________ (подпись)")
    p = doc.add_paragraph()
    p.add_run(f"Научный руководитель ВКР, {args.advisor_title} {args.advisor}    ___________________ (подпись)")
    p = doc.add_paragraph()
    p.add_run(f"Нормоконтролёр       {args.normocontroller or '__________________'}    ___________________ (подпись)")
    add("")
    add("")
    add(f"Самара, {args.year}")

    # Восстановим sectPr в конце body
    if sectPr is not None:
        body.append(sectPr)

    out_path = Path(args.output).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"OK: написан {out_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Сборка titlepage-filled.docx из метаданных студента")
    parser.add_argument("--template", default="template/blank-template.docx")
    parser.add_argument("--output", default="template/titlepage-filled.docx")
    parser.add_argument("--topic", required=True)
    parser.add_argument("--student", required=True, help="ФИО как 'И.И. Иванов'")
    parser.add_argument("--group", default="")
    parser.add_argument("--direction", required=True, help="01.03.02 Прикладная математика и информатика, профиль ...")
    parser.add_argument("--advisor", required=True, help="ФИО научрука как 'А.А. Иванов'")
    parser.add_argument("--advisor-title", default="к.ф.-м.н., доцент", help="Степень и должность научрука")
    parser.add_argument("--institute", required=True, default="Институт информатики и кибернетики")
    parser.add_argument("--department", required=True, default="Кафедра технической кибернетики")
    parser.add_argument("--university", default="Самарский национальный исследовательский университет имени академика С.П. Королёва")
    parser.add_argument("--normocontroller", default="")
    parser.add_argument("--year", required=True)

    args = parser.parse_args()
    build_titlepage(args)


if __name__ == "__main__":
    main()
